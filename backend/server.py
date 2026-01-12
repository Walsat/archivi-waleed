from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime
from bson import ObjectId
import base64
import io
from PIL import Image
import PyPDF2
import docx

# Import AI libraries
from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Get Emergent LLM Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============ Models ============

class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    username: str
    password: str
    full_name: str
    role: str = "موظف"
    created_at: datetime = Field(default_factory=datetime.utcnow)

class UserCreate(BaseModel):
    username: str
    password: str
    full_name: str
    role: Optional[str] = "موظف"

class UserLogin(BaseModel):
    username: str
    password: str

class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: Optional[str] = ""
    file_type: str  # image, pdf, word
    file_data: str  # base64 encoded
    category: Optional[str] = ""
    owner_name: Optional[str] = ""
    land_type: Optional[str] = ""
    location: Optional[str] = ""
    extracted_text: Optional[str] = ""
    summary: Optional[str] = ""
    auto_category: Optional[str] = ""
    keywords: Optional[List[str]] = []
    notes: Optional[str] = ""
    uploaded_by: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class DocumentCreate(BaseModel):
    title: str
    description: Optional[str] = ""
    file_type: str
    file_data: str
    category: Optional[str] = ""
    owner_name: Optional[str] = ""
    land_type: Optional[str] = ""
    location: Optional[str] = ""
    notes: Optional[str] = ""
    uploaded_by: str

class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    category: Optional[str] = None
    owner_name: Optional[str] = None
    land_type: Optional[str] = None
    location: Optional[str] = None
    notes: Optional[str] = None

class SearchQuery(BaseModel):
    query: str
    category: Optional[str] = None
    land_type: Optional[str] = None
    date_from: Optional[str] = None
    date_to: Optional[str] = None


# ============ Helper Functions ============

def extract_text_from_image(base64_data: str) -> str:
    """Extract text from image using base64 data"""
    return base64_data

def extract_text_from_pdf(base64_data: str) -> str:
    """Extract text from PDF"""
    try:
        pdf_bytes = base64.b64decode(base64_data)
        pdf_file = io.BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(base64_data: str) -> str:
    """Extract text from Word document"""
    try:
        doc_bytes = base64.b64decode(base64_data)
        doc_file = io.BytesIO(doc_bytes)
        doc = docx.Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from Word: {e}")
        return ""

async def process_document_with_ai(file_data: str, file_type: str, title: str) -> dict:
    """Process document with AI for OCR, classification, and summary"""
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message="أنت مساعد ذكي متخصص في تحليل الوثائق الزراعية العراقية. مهمتك استخراج النصوص وتصنيف الوثائق وإنشاء ملخصات باللغة العربية."
        ).with_model("openai", "gpt-5.2")
        
        extracted_text = ""
        
        # Extract text based on file type
        if file_type == "image":
            # Use Vision API for OCR
            image_content = ImageContent(image_base64=file_data)
            message = UserMessage(
                text="استخرج كل النصوص العربية والإنجليزية من هذه الصورة. اكتب النص كما هو بدقة.",
                file_contents=[image_content]
            )
            extracted_text = await chat.send_message(message)
            
        elif file_type == "pdf":
            extracted_text = extract_text_from_pdf(file_data)
            
        elif file_type == "word":
            extracted_text = extract_text_from_docx(file_data)
        
        # If no text extracted, return empty result
        if not extracted_text or len(extracted_text.strip()) < 10:
            return {
                "extracted_text": "",
                "summary": "",
                "auto_category": "غير مصنف",
                "keywords": []
            }
        
        # Create new chat for classification and summary
        analysis_chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message="أنت محلل وثائق متخصص في الأراضي الزراعية."
        ).with_model("openai", "gpt-5.2")
        
        # Classify and summarize
        analysis_prompt = f"""
لديك وثيقة بعنوان: {title}

النص المستخرج:
{extracted_text[:2000]}

المطلوب:
1. صنف هذه الوثيقة إلى إحدى الفئات التالية فقط: (سند ملكية، عقد إيجار، خريطة مساحية، تقرير فني، طلب خدمة، شهادة، أخرى)
2. اكتب ملخص مختصر للوثيقة (2-3 جمل)
3. استخرج الكلمات المفتاحية المهمة (5-7 كلمات)

الرجاء الرد بالصيغة التالية فقط:
التصنيف: [الفئة]
الملخص: [الملخص]
الكلمات المفتاحية: [كلمة1، كلمة2، كلمة3]
"""
        
        analysis_result = await analysis_chat.send_message(UserMessage(text=analysis_prompt))
        
        # Parse the result
        category = "أخرى"
        summary = ""
        keywords = []
        
        if analysis_result:
            lines = analysis_result.split('\n')
            for line in lines:
                if line.startswith('التصنيف:'):
                    category = line.replace('التصنيف:', '').strip()
                elif line.startswith('الملخص:'):
                    summary = line.replace('الملخص:', '').strip()
                elif line.startswith('الكلمات المفتاحية:'):
                    keywords_str = line.replace('الكلمات المفتاحية:', '').strip()
                    keywords = [k.strip() for k in keywords_str.split('،')]
        
        return {
            "extracted_text": extracted_text[:5000],  # Limit size
            "summary": summary,
            "auto_category": category,
            "keywords": keywords
        }
        
    except Exception as e:
        logger.error(f"Error processing document with AI: {e}")
        return {
            "extracted_text": "",
            "summary": "",
            "auto_category": "خطأ في التصنيف",
            "keywords": []
        }


async def smart_search(query: str, filters: dict) -> List[dict]:
    """Smart search using AI to understand Arabic queries"""
    try:
        # Build MongoDB query
        mongo_query = {}
        
        # Add filters
        if filters.get('category'):
            mongo_query['category'] = filters['category']
        if filters.get('land_type'):
            mongo_query['land_type'] = filters['land_type']
        
        # Search in multiple fields
        if query:
            mongo_query['$or'] = [
                {'title': {'$regex': query, '$options': 'i'}},
                {'description': {'$regex': query, '$options': 'i'}},
                {'extracted_text': {'$regex': query, '$options': 'i'}},
                {'owner_name': {'$regex': query, '$options': 'i'}},
                {'location': {'$regex': query, '$options': 'i'}},
                {'keywords': {'$regex': query, '$options': 'i'}},
            ]
        
        # Execute search with projection to exclude large file_data field
        documents = await db.documents.find(
            mongo_query,
            {'file_data': 0}  # Exclude base64 file data for performance
        ).sort('created_at', -1).limit(100).to_list(100)
        
        return documents
        
    except Exception as e:
        logger.error(f"Error in smart search: {e}")
        return []


# ============ Auth Routes ============

@api_router.post("/auth/register", response_model=User)
async def register(user_data: UserCreate):
    """Register new user"""
    # Check if username exists
    existing_user = await db.users.find_one({"username": user_data.username})
    if existing_user:
        raise HTTPException(status_code=400, detail="اسم المستخدم موجود بالفعل")
    
    # Create user (in production, hash the password)
    user = User(**user_data.dict())
    await db.users.insert_one(user.dict())
    
    return user

@api_router.post("/auth/login")
async def login(credentials: UserLogin):
    """Login user"""
    user = await db.users.find_one({"username": credentials.username})
    
    if not user or user['password'] != credentials.password:
        raise HTTPException(status_code=401, detail="اسم المستخدم أو كلمة المرور غير صحيحة")
    
    return {
        "success": True,
        "user": {
            "id": user['id'],
            "username": user['username'],
            "full_name": user['full_name'],
            "role": user['role']
        }
    }


# ============ Document Routes ============

@api_router.post("/documents", response_model=Document)
async def create_document(doc_data: DocumentCreate):
    """Upload and process new document"""
    try:
        # Process with AI
        ai_result = await process_document_with_ai(
            doc_data.file_data,
            doc_data.file_type,
            doc_data.title
        )
        
        # Create document
        document = Document(
            **doc_data.dict(),
            extracted_text=ai_result['extracted_text'],
            summary=ai_result['summary'],
            auto_category=ai_result['auto_category'],
            keywords=ai_result['keywords']
        )
        
        await db.documents.insert_one(document.dict())
        
        return document
        
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise HTTPException(status_code=500, detail=f"خطأ في رفع الوثيقة: {str(e)}")

@api_router.get("/documents", response_model=List[Document])
async def get_documents(
    category: Optional[str] = None,
    land_type: Optional[str] = None,
    limit: int = 50
):
    """Get all documents with optional filters"""
    query = {}
    
    if category:
        query['category'] = category
    if land_type:
        query['land_type'] = land_type
    
    documents = await db.documents.find(query).sort('created_at', -1).limit(limit).to_list(limit)
    return [Document(**doc) for doc in documents]

@api_router.get("/documents/{document_id}", response_model=Document)
async def get_document(document_id: str):
    """Get single document by ID"""
    document = await db.documents.find_one({"id": document_id})
    
    if not document:
        raise HTTPException(status_code=404, detail="الوثيقة غير موجودة")
    
    return Document(**document)

@api_router.put("/documents/{document_id}", response_model=Document)
async def update_document(document_id: str, update_data: DocumentUpdate):
    """Update document details"""
    # Get existing document
    existing_doc = await db.documents.find_one({"id": document_id})
    if not existing_doc:
        raise HTTPException(status_code=404, detail="الوثيقة غير موجودة")
    
    # Update fields
    update_dict = {k: v for k, v in update_data.dict().items() if v is not None}
    update_dict['updated_at'] = datetime.utcnow()
    
    await db.documents.update_one(
        {"id": document_id},
        {"$set": update_dict}
    )
    
    # Get updated document
    updated_doc = await db.documents.find_one({"id": document_id})
    return Document(**updated_doc)

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete document"""
    result = await db.documents.delete_one({"id": document_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="الوثيقة غير موجودة")
    
    return {"success": True, "message": "تم حذف الوثيقة بنجاح"}

@api_router.post("/documents/search")
async def search_documents(search_query: SearchQuery):
    """Smart search in documents"""
    try:
        filters = {}
        if search_query.category:
            filters['category'] = search_query.category
        if search_query.land_type:
            filters['land_type'] = search_query.land_type
        
        results = await smart_search(search_query.query, filters)
        
        return {
            "success": True,
            "results": [Document(**doc) for doc in results],
            "count": len(results)
        }
        
    except Exception as e:
        logger.error(f"Error in search: {e}")
        raise HTTPException(status_code=500, detail=f"خطأ في البحث: {str(e)}")


# ============ Statistics Routes ============

@api_router.get("/stats")
async def get_statistics():
    """Get system statistics"""
    try:
        total_documents = await db.documents.count_documents({})
        total_users = await db.users.count_documents({})
        
        # Get documents by category
        pipeline = [
            {"$group": {"_id": "$auto_category", "count": {"$sum": 1}}}
        ]
        categories = await db.documents.aggregate(pipeline).to_list(100)
        
        # Get recent documents
        recent_docs = await db.documents.find().sort('created_at', -1).limit(5).to_list(5)
        
        return {
            "total_documents": total_documents,
            "total_users": total_users,
            "by_category": {cat['_id']: cat['count'] for cat in categories},
            "recent_documents": [
                {
                    "id": doc['id'],
                    "title": doc['title'],
                    "created_at": doc['created_at']
                }
                for doc in recent_docs
            ]
        }
        
    except Exception as e:
        logger.error(f"Error getting statistics: {e}")
        raise HTTPException(status_code=500, detail=f"خطأ في الحصول على الإحصائيات: {str(e)}")


# ============ Test Routes ============

@api_router.get("/")
async def root():
    return {
        "message": "نظام الأرشفة الإلكترونية - مديرية زراعة صلاح الدين",
        "version": "1.0.0",
        "status": "running"
    }

@api_router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.utcnow()}


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
